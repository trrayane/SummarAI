"""
routes.py
Routes FastAPI — validation, streaming SSE, export DOCX.
"""

import io
import uuid
import logging

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import StreamingResponse, Response

from agents.summarizer_agent import SummarizerAgent, LANG_REJECTED_TOKEN
from services.history_service import HistoryService
from services.gemini_service import gemini_service
from memory.session_store import save_document, get_document
from tools.web_scraper import extract_text_from_url

logger = logging.getLogger(__name__)

router = APIRouter()
agent  = SummarizerAgent()

MAX_FILE_SIZE   = 10 * 1024 * 1024
MAX_TEXT_LENGTH = 100_000
ALLOWED_MIME    = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}
ALLOWED_EXT    = {".pdf", ".docx", ".txt"}
ALLOWED_STYLES = {"concis", "détaillé", "bullet"}


class FileWrapper:
    def __init__(self, filename: str, contents: bytes):
        self.filename = filename
        self.file     = io.BytesIO(contents)


def _sse_encode(token: str) -> str:
    return token.replace("\n", "\\n")


# ---------------------------------------------------------------------------
# POST /summarize
# ---------------------------------------------------------------------------

@router.post("/summarize")
async def summarize(
    text:  str        = Form(default=""),
    url:   str        = Form(default=""),
    style: str        = Form(default="concis"),
    file:  UploadFile = File(default=None),
):
    if style not in ALLOWED_STYLES:
        style = "concis"

    file_wrapper = None
    source_text  = text

    if text and len(text) > MAX_TEXT_LENGTH:
        raise HTTPException(400, f"Text too long (max {MAX_TEXT_LENGTH} characters).")

    if url and url.strip():
        try:
            source_text = await extract_text_from_url(url.strip())
        except (ValueError, RuntimeError) as exc:
            async def _err():
                yield f"data: ⚠️ Unable to load URL: {exc}\n\n"
                yield "data: [DONE]\n\n"
            return StreamingResponse(_err(), media_type="text/event-stream")

    elif file and file.filename:
        ext = "." + file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
        if ext not in ALLOWED_EXT:
            raise HTTPException(415, f"Unsupported format: {ext}. Accepted: {', '.join(ALLOWED_EXT)}")

        contents = await file.read()
        if len(contents) > MAX_FILE_SIZE:
            raise HTTPException(413, f"File too large (max {MAX_FILE_SIZE // (1024*1024)} MB).")

        file_wrapper = FileWrapper(filename=file.filename, contents=contents)
        source_text  = ""

    session_id = str(uuid.uuid4())

    async def stream_response():
        # Premier chunk : session_id (temporaire, peut être annulé si langue rejetée)
        yield f"data: __SESSION__{session_id}__\n\n"

        full_summary  = ""
        lang_rejected = False   # ← flag de rejet de langue

        async for token in agent.summarize(text=source_text, style=style, file=file_wrapper):

            # ── Détection du token de rejet langue ──────────────────────
            if token == LANG_REJECTED_TOKEN:
                lang_rejected = True
                # On signale au frontend de désactiver le Q&A
                yield "data: __NO_QA__\n\n"
                continue   # on n'ajoute pas ce token au résumé

            full_summary += token
            yield f"data: {_sse_encode(token)}\n\n"

        yield "data: [DONE]\n\n"

        # ── Stockage session : UNIQUEMENT si la langue est anglaise ─────
        if not lang_rejected:
            save_document(session_id, source_text or full_summary)
        else:
            logger.info(
                "Session %s non sauvegardée — texte non-anglais rejeté.",
                session_id
            )

    return StreamingResponse(stream_response(), media_type="text/event-stream")


# ---------------------------------------------------------------------------
# POST /ask
# ---------------------------------------------------------------------------

@router.post("/ask")
async def ask_question(
    session_id: str = Form(...),
    question:   str = Form(...),
):
    """Q&A sur le document original lié au session_id."""
    document = get_document(session_id)
    if not document:
        raise HTTPException(
            404,
            "Session not found or expired. Please regenerate a summary first."
        )
    if not question or not question.strip():
        raise HTTPException(400, "Question cannot be empty.")

    async def qa_stream():
        async for token in agent.answer_question(
            document=document,
            question=question.strip()
        ):
            yield f"data: {_sse_encode(token)}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(qa_stream(), media_type="text/event-stream")


# ---------------------------------------------------------------------------
# GET /history
# ---------------------------------------------------------------------------

@router.get("/history")
async def get_history(limit: int = 10):
    service = HistoryService()
    return await service.get_recent(limit)


# ---------------------------------------------------------------------------
# POST /export/docx
# ---------------------------------------------------------------------------

@router.post("/export/docx")
async def export_docx(
    title:   str = Form(...),
    summary: str = Form(...),
    style:   str = Form(default=""),
    source:  str = Form(default=""),
):
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "python-docx not installed.")

    doc = DocxDocument()
    doc.core_properties.author = "SummAI"
    doc.core_properties.title  = title

    heading           = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run               = heading.runs[0]
    run.font.size     = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x18)

    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Style: {style}   |   Source: {source}").font.size = Pt(9)
    meta_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x80)
    doc.add_paragraph()

    lines = summary.split("\n")
    is_bullet = any(l.strip().startswith("•") for l in lines)

    if is_bullet:
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("•"):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(stripped.lstrip("• ").strip())
            elif stripped:
                doc.add_paragraph(stripped)
    else:
        for paragraph in summary.split("\n\n"):
            p   = doc.add_paragraph()
            run = p.add_run(paragraph.strip())
            run.font.size = Pt(11)

    doc.add_paragraph()
    footer = doc.add_paragraph("Generated by SummAI")
    footer.runs[0].font.size      = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0xB0, 0xB0, 0xA8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)[:40]
    filename   = f"SummAI_{safe_title}.docx"

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# GET /health
# ---------------------------------------------------------------------------

@router.get("/health")
async def health():
    gemini_status = await gemini_service.health_check()
    return {
        "status": "ok",
        "gemini": gemini_status,
        "cache_size": gemini_service.cache_size(),
    }